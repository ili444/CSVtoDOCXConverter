# -*- coding: utf-8 -*-
# coding: utf-8
import csv
import os
import docx

def csv_reader(file_obj):
    """
    Read a csv file
    """
    reader = csv.reader(file_obj)
    csv_read = [iters for iters in reader]
    date = []
    for iters in csv_read:
        for y in iters:
            a = y.replace('"', '').split(';')
            date.append([a[16], a[17], a[19], a[18], a[20]])
    return date, len(date), len(date[0])

def record_docx(date, rows, cols, doc_file):
    doc = docx.Document()

    par = doc.add_paragraph()
    par.add_run('С инструктажем по технике безопасности ознакомлен, правила мне ясны. '
                      'Я понимаю возможные наступления последствий в виде травм'
                      ' в связи с неисполнением ТБ и указаний гида/ экскурсовода/сотрудника «Центра путешественников» во время мероприятия__________________ дата____________\n').bold=True

    # добавляем таблицу NxM
    table = doc.add_table(rows=rows, cols=cols+1)
    # применяем стиль для таблицы
    table.style = 'Table Grid'

    # заполняем таблицу данными
    for row in range(rows):
        for col in range(cols):
            # получаем ячейку таблицы
            cell = table.cell(row, col)
            # записываем в ячейку данные
            cell.text = date[row][col]

    cell = table.cell(0, cols)
    cell.text = 'Подпись'

    doc.save(doc_file)
    return 'ok'



if __name__ == "__main__":
    csv_path = os.path.abspath(input('Введите название csv-файла:') + '.csv')
    doc_file = input('Введите название docx-файла на выходе:') + '.docx'
    doc_path = os.path.abspath(doc_file)
    with open(csv_path, "r", encoding='utf-8') as f_obj:
        date, rows, cols = csv_reader(f_obj)
        record_docx(date, rows, cols, doc_file)
